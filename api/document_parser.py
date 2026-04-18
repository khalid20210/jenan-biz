"""
document_parser.py — مُشرّح المستندات العربية (PDF & DOCX)
يستخرج النص مع الحفاظ على التسلسل الهرمي ويقسّمه حسب المواد القانونية
"""

import re
import io
from typing import List, Dict, Any

# ─── استيراد مكتبات القراءة ──────────────────────────────────────
try:
    import pdfplumber
    _PDF_PARSE = True
except ImportError:
    _PDF_PARSE = False

try:
    from docx import Document as DocxDocument
    _DOCX_PARSE = True
except ImportError:
    _DOCX_PARSE = False


# ─── الأنماط العربية لتقسيم المواد ──────────────────────────────
_ARTICLE_PATTERNS = [
    # المادة الأولى / المادة 1 / المادة (1)
    r'(?:^|\n)\s*(المادة\s+(?:الأولى|الثانية|الثالثة|الرابعة|الخامسة|السادسة|السابعة|الثامنة|التاسعة|العاشرة|\d+|(?:الحادية|الثانية|الثالثة)\s+عشرة?))',
    # الفصل الأول / الفصل 1
    r'(?:^|\n)\s*(الفصل\s+(?:الأول|الثاني|الثالث|الرابع|الخامس|\d+))',
    # الباب الأول
    r'(?:^|\n)\s*(الباب\s+(?:الأول|الثاني|الثالث|الرابع|\d+))',
    # البند أولاً / ثانياً
    r'(?:^|\n)\s*((?:أولاً|ثانياً|ثالثاً|رابعاً|خامساً|سادساً|سابعاً|ثامناً|تاسعاً|عاشراً)\s*[:\-])',
]

_COMPILED_PATTERNS = [re.compile(p, re.MULTILINE | re.UNICODE) for p in _ARTICLE_PATTERNS]

# أنماط العناوين في DOCX
_HEADING_STYLES = {'heading 1', 'heading 2', 'heading 3', 'عنوان 1', 'عنوان 2', 'عنوان 3'}


def _is_article_header(text: str) -> bool:
    """يتحقق إذا كان النص بداية مادة/فصل/باب."""
    stripped = text.strip()
    return any(p.match(stripped) for p in _COMPILED_PATTERNS)


def _split_into_chunks(full_text: str, source_name: str) -> List[Dict[str, Any]]:
    """
    يقسّم النص الكامل إلى chunks حسب المواد.
    إذا لم يجد مواد → يقسّم بحد أقصى 800 حرف مع تداخل 100 حرف.
    """
    chunks: List[Dict[str, Any]] = []

    # محاولة التقسيم الدلالي حسب المواد
    all_positions = []
    for pattern in _COMPILED_PATTERNS:
        for m in pattern.finditer(full_text):
            all_positions.append(m.start())

    all_positions = sorted(set(all_positions))

    if len(all_positions) >= 2:
        # تقسيم عند كل مادة
        boundaries = all_positions + [len(full_text)]
        for i in range(len(all_positions)):
            start = all_positions[i]
            end = boundaries[i + 1]
            chunk_text = full_text[start:end].strip()
            if len(chunk_text) < 20:
                continue
            # استخراج عنوان المادة
            first_line = chunk_text.split('\n')[0].strip()
            chunks.append({
                'chunk_id': i + 1,
                'header': first_line[:120],
                'text': chunk_text,
                'char_count': len(chunk_text),
                'source': source_name,
            })
    else:
        # Fallback: تقسيم بـ 800 حرف مع تداخل 100 حرف
        MAX = 800
        OVERLAP = 100
        pos = 0
        idx = 1
        while pos < len(full_text):
            chunk_text = full_text[pos:pos + MAX].strip()
            if chunk_text:
                chunks.append({
                    'chunk_id': idx,
                    'header': f'قطعة {idx}',
                    'text': chunk_text,
                    'char_count': len(chunk_text),
                    'source': source_name,
                })
                idx += 1
            pos += MAX - OVERLAP

    return chunks


# ─── استخراج PDF ─────────────────────────────────────────────────

def parse_pdf(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    يستخرج النص من PDF عربي مع الحفاظ على الترتيب.
    يستخدم pdfplumber الذي يدعم RTL وترتيب الكلمات العربية.
    """
    if not _PDF_PARSE:
        raise RuntimeError("مكتبة pdfplumber غير مثبّتة")

    pages_text: List[str] = []
    metadata: Dict[str, Any] = {}

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        metadata = {
            'pages': len(pdf.pages),
            'filename': filename,
        }
        for page in pdf.pages:
            # استخراج النص مع الحفاظ على ترتيب الكلمات
            text = page.extract_text(x_tolerance=3, y_tolerance=3) or ''
            # تنظيف: إزالة أسطر فارغة متعددة
            text = re.sub(r'\n{3,}', '\n\n', text)
            if text.strip():
                pages_text.append(text)

    full_text = '\n\n'.join(pages_text)
    # إزالة المسافات الزائدة في بداية الأسطر (شائعة في PDF العربي)
    full_text = re.sub(r'[ \t]+\n', '\n', full_text)
    full_text = re.sub(r'\n[ \t]+', '\n', full_text)

    chunks = _split_into_chunks(full_text, filename)

    return {
        'type': 'pdf',
        'filename': filename,
        'metadata': metadata,
        'full_text': full_text,
        'chunks': chunks,
        'total_chunks': len(chunks),
        'total_chars': len(full_text),
    }


# ─── استخراج DOCX ────────────────────────────────────────────────

def parse_docx(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    يستخرج النص من DOCX مع الحفاظ على التسلسل الهرمي للعناوين.
    يُميّز بين الـ headings والفقرات والقوائم.
    """
    if not _DOCX_PARSE:
        raise RuntimeError("مكتبة python-docx غير مثبّتة")

    doc = DocxDocument(io.BytesIO(file_bytes))
    structured: List[Dict[str, str]] = []
    full_lines: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name.lower() if para.style else ''
        is_heading = any(h in style_name for h in _HEADING_STYLES)

        block = {
            'type': 'heading' if is_heading else 'paragraph',
            'style': para.style.name if para.style else 'Normal',
            'text': text,
        }
        structured.append(block)

        # للنص الكامل: أضف سطراً فارغاً بعد العناوين
        if is_heading:
            full_lines.append(f'\n{text}\n')
        else:
            full_lines.append(text)

    full_text = '\n'.join(full_lines)
    full_text = re.sub(r'\n{3,}', '\n\n', full_text)

    chunks = _split_into_chunks(full_text, filename)

    return {
        'type': 'docx',
        'filename': filename,
        'metadata': {
            'paragraphs': len(doc.paragraphs),
            'filename': filename,
        },
        'structured': structured,
        'full_text': full_text,
        'chunks': chunks,
        'total_chunks': len(chunks),
        'total_chars': len(full_text),
    }


# ─── دالة موحّدة ─────────────────────────────────────────────────

def parse_document(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """نقطة دخول واحدة — تكتشف نوع الملف تلقائياً."""
    name_lower = filename.lower()
    if name_lower.endswith('.pdf'):
        return parse_pdf(file_bytes, filename)
    elif name_lower.endswith('.docx'):
        return parse_docx(file_bytes, filename)
    else:
        raise ValueError(f"نوع الملف غير مدعوم: {filename} — يُدعم فقط PDF و DOCX")
